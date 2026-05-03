"""
File parsing service.
Handles: PDF, DOCX, TXT, images, CSV, ZIP, PY, JS, HTML, CSS, JSON, YAML, and more.
"""
import os
import base64
import io
import zipfile
from pathlib import Path
from typing import Tuple

# Max chars extracted per file
MAX_CHARS = 12000
MAX_ZIP_FILES = 20  # max files to read inside a zip


def parse_file(filename: str, file_bytes: bytes) -> Tuple[str, str, str]:
    """
    Returns: (extracted_text, image_b64_or_empty, mime_type)
    """
    ext = Path(filename).suffix.lower()

    # ── IMAGES ──────────────────────────────────────────────────────
    if ext in (".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp", ".svg"):
        if ext == ".svg":
            # SVG is text
            return file_bytes.decode("utf-8", errors="replace")[:MAX_CHARS], "", "image/svg+xml"
        b64 = base64.b64encode(file_bytes).decode()
        mime = {".jpg":"image/jpeg",".jpeg":"image/jpeg",".png":"image/png",
                ".webp":"image/webp",".gif":"image/gif",".bmp":"image/bmp"}.get(ext,"image/jpeg")
        return "", b64, mime

    # ── PDF ─────────────────────────────────────────────────────────
    if ext == ".pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            pages = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    pages.append(f"[Page {i+1}]\n{text.strip()}")
            extracted = "\n\n".join(pages)
            if not extracted.strip():
                return "[PDF has no extractable text — may be scanned image]", "", "application/pdf"
            return extracted[:MAX_CHARS], "", "application/pdf"
        except ImportError:
            return "[pypdf not installed]", "", "application/pdf"
        except Exception as e:
            return f"[PDF error: {e}]", "", "application/pdf"

    # ── DOCX ────────────────────────────────────────────────────────
    if ext in (".docx", ".doc"):
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        parts.append(row_text)
            return "\n".join(parts)[:MAX_CHARS], "", "application/docx"
        except ImportError:
            return "[python-docx not installed]", "", "application/docx"
        except Exception as e:
            return f"[DOCX error: {e}]", "", "application/docx"

    # ── ZIP ─────────────────────────────────────────────────────────
    if ext == ".zip":
        try:
            results = []
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
                all_files = zf.namelist()
                results.append(f"ZIP archive containing {len(all_files)} files:\n")
                results.append("\n".join(f"  📄 {f}" for f in all_files[:50]))
                if len(all_files) > 50:
                    results.append(f"  ... and {len(all_files)-50} more files")
                results.append("\n\n--- File Contents ---\n")

                readable_count = 0
                for name in all_files:
                    if readable_count >= MAX_ZIP_FILES:
                        results.append(f"\n[Stopped at {MAX_ZIP_FILES} files — too many to show]")
                        break
                    # Skip directories and binary files
                    if name.endswith("/"):
                        continue
                    fext = Path(name).suffix.lower()
                    if fext in (".jpg",".jpeg",".png",".gif",".webp",".bmp",
                                ".exe",".dll",".so",".bin",".pyc",".class",
                                ".zip",".tar",".gz",".rar",".7z",".mp4",
                                ".mp3",".wav",".avi",".mov"):
                        results.append(f"\n[{name}] — binary file, skipped")
                        continue
                    try:
                        content = zf.read(name).decode("utf-8", errors="replace")
                        if len(content) > 3000:
                            content = content[:3000] + f"\n... [{len(content)-3000} more chars]"
                        results.append(f"\n{'='*40}\n📄 {name}\n{'='*40}\n{content}")
                        readable_count += 1
                    except Exception:
                        results.append(f"\n[{name}] — could not read")

            return "\n".join(results)[:MAX_CHARS], "", "application/zip"
        except zipfile.BadZipFile:
            return "[Invalid ZIP file]", "", "application/zip"
        except Exception as e:
            return f"[ZIP error: {e}]", "", "application/zip"

    # ── CSV ─────────────────────────────────────────────────────────
    if ext == ".csv":
        try:
            import csv
            text = file_bytes.decode("utf-8", errors="replace")
            rows = list(csv.reader(io.StringIO(text)))
            lines = [" | ".join(row) for row in rows[:100]]
            result = f"CSV file ({len(rows)} rows, {len(rows[0]) if rows else 0} columns):\n" + "\n".join(lines)
            if len(rows) > 100:
                result += f"\n... and {len(rows)-100} more rows"
            return result[:MAX_CHARS], "", "text/csv"
        except Exception as e:
            return f"[CSV error: {e}]", "", "text/csv"

    # ── EXCEL ────────────────────────────────────────────────────────
    if ext in (".xlsx", ".xls"):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
            parts = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                parts.append(f"\n[Sheet: {sheet_name}]")
                row_count = 0
                for row in ws.iter_rows(values_only=True):
                    if row_count >= 100:
                        parts.append("... [truncated]")
                        break
                    row_str = " | ".join(str(c) if c is not None else "" for c in row)
                    if row_str.strip():
                        parts.append(row_str)
                        row_count += 1
            return "\n".join(parts)[:MAX_CHARS], "", "application/xlsx"
        except ImportError:
            # Fallback without openpyxl
            return "[Excel reading requires openpyxl — pip install openpyxl]", "", "application/xlsx"
        except Exception as e:
            return f"[Excel error: {e}]", "", "application/xlsx"

    # ── ALL TEXT / CODE FILES ────────────────────────────────────────
    text_extensions = {
        # Web
        ".html", ".htm", ".css", ".js", ".jsx", ".ts", ".tsx", ".vue", ".svelte",
        # Python
        ".py", ".pyw", ".ipynb",
        # Data
        ".json", ".yaml", ".yml", ".toml", ".xml", ".env", ".ini", ".cfg", ".conf",
        # Docs
        ".txt", ".md", ".markdown", ".rst", ".log",
        # Systems / backend
        ".sh", ".bash", ".zsh", ".fish", ".ps1", ".bat", ".cmd",
        ".rs", ".go", ".java", ".kt", ".c", ".cpp", ".h", ".hpp",
        ".cs", ".rb", ".php", ".swift", ".r", ".sql", ".lua",
        # Config / infra
        ".dockerfile", ".tf", ".hcl", ".gradle", ".makefile",
        # Other
        ".gitignore", ".htaccess", ".editorconfig",
    }

    # Also match files with no extension or Dockerfile/Makefile etc
    base = Path(filename).name.lower()
    no_ext_readable = base in ("dockerfile", "makefile", "procfile", "requirements",
                                "gemfile", "rakefile", "vagrantfile", "jenkinsfile",
                                "readme", "license", "changelog", "authors")

    if ext in text_extensions or no_ext_readable or ext == "":
        try:
            text = file_bytes.decode("utf-8", errors="replace")
            lines = text.split("\n")
            line_count = len(lines)
            if len(text) > MAX_CHARS:
                text = text[:MAX_CHARS] + f"\n\n... [{len(text)-MAX_CHARS} more characters — file truncated]"
            header = f"[{filename}] — {line_count} lines, {len(file_bytes)} bytes\n\n"
            return header + text, "", "text/plain"
        except Exception as e:
            return f"[Read error: {e}]", "", "text/plain"

    # ── UNKNOWN / BINARY ────────────────────────────────────────────
    return (
        f"[Unsupported file type: {ext}]\n"
        f"File: {filename} ({len(file_bytes)} bytes)\n"
        f"Supported: PDF, DOCX, ZIP, CSV, XLSX, PY, JS, TS, HTML, CSS, JSON, YAML, TXT, MD, and all code files.",
        "", "application/octet-stream"
    )


def format_file_context(filename: str, text: str) -> str:
    """Wrap extracted file content into a clear prompt context block."""
    return (
        f'The user uploaded the file: "{filename}"\n\n'
        f'File contents:\n'
        f'```\n{text}\n```\n\n'
        f'Analyze this file and answer the user\'s question about it.'
                        )
    
